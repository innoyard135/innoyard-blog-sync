"""PipelineResult → Word(.docx)."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

from src.models import ContentIdea, PipelineResult


_EMPHASIS_RE = re.compile(r"(\*\*|__)(.+?)\1")


def _strip_emphasis(text: str) -> str:
    """마크다운 굵게/이탤릭 강조 표시(**text**, __text__)는 텍스트만 남기고 제거."""
    prev = None
    while prev != text:
        prev = text
        text = _EMPHASIS_RE.sub(r"\2", text)
    return text


def _add_markdownish_body(doc: Document, body: str) -> None:
    for line in body.splitlines():
        line = _strip_emphasis(line.rstrip())
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)


def _add_ideas_section(doc: Document, ideas: list[ContentIdea]) -> None:
    doc.add_page_break()
    h = doc.add_heading("추가 콘텐츠 아이템 제안", level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    intro = doc.add_paragraph(
        "일반 부모가 검색·블로그 주제로 잘 떠올리지 못하는 각도입니다. "
        "원고와 연계해 시리즈·카드뉴스·숏폼으로 확장할 수 있습니다."
    )
    intro.runs[0].font.size = Pt(10)
    intro.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    for i, idea in enumerate(ideas, start=1):
        doc.add_heading(f"{i}. {idea.topic}", level=2)
        for label, value in [
            ("훅", idea.hook),
            ("전문가 각도", idea.expert_angle),
            ("왜 놓치기 쉬운가", idea.why_missed),
        ]:
            if value:
                p = doc.add_paragraph()
                run_label = p.add_run(f"{label}: ")
                run_label.bold = True
                p.add_run(value)


def export_docx(result: PipelineResult, out_path: Path, *, author_label: str) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    doc.add_heading(result.title, level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"작성 관점: {author_label}").italic = True
    if result.tags:
        doc.add_paragraph("태그: " + ", ".join(result.tags))

    _add_markdownish_body(doc, result.body)

    disclaimer = doc.add_paragraph()
    disclaimer.add_run(
        "\n본 글은 일반적인 건강 정보이며, 진단·치료는 직접 진료를 통해 결정하시기 바랍니다."
    ).italic = True

    if result.ideas:
        _add_ideas_section(doc, result.ideas)

    if result.source_path:
        foot = doc.add_paragraph()
        foot.add_run(f"\n원고 출처: {result.source_path}").font.size = Pt(9)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def safe_filename(title: str) -> str:
    s = re.sub(r'[\\/:*?"<>|]', "_", title)
    return (s[:60] or "untitled").strip()
